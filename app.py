import streamlit as st
import google.generativeai as genai
from docx import Document
import io
import os

# Page configuration
st.set_page_config(page_title="AI PDF to Word Converter", page_icon="📄", layout="centered")

# Custom CSS for a professional look
st.markdown("""
    <style>
    .main {
        background-color: #f8f9fa;
    }
    .stButton>button {
        width: 100%;
        border-radius: 20px;
        height: 3em;
        background-color: #007bff;
        color: white;
    }
    .stDownloadButton>button {
        width: 100%;
        border-radius: 20px;
        height: 3em;
        background-color: #28a745;
        color: white;
    }
    </style>
    """, unsafe_allow_html=True)

def markdown_to_docx(markdown_text):
    doc = Document()
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line:
            doc.add_paragraph(line)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def main():
    st.title("📄 AI PDF to Word Converter")
    st.write("Convert your PDF documents to high-quality Word files using Gemini 1.5 Flash.")

    # API Key handling
    api_key = st.secrets.get("GEMINI_API_KEY")
    if not api_key:
        api_key = st.sidebar.text_input("Enter Gemini API Key", type="password")
    
    if not api_key:
        st.warning("Please provide a Gemini API Key to proceed.")
        st.info("You can get one at https://aistudio.google.com/app/apikey")
        return

    genai.configure(api_key=api_key)

    uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")

    if uploaded_file is not None:
        st.info(f"File uploaded: {uploaded_file.name}")
        
        if st.button("Convert to Word"):
            try:
                with st.spinner("AI is analyzing the PDF..."):
                    progress_bar = st.progress(0)
                    
                    # Upload file to Gemini
                    # Note: For Streamlit, we convert to bytes and send as part of the prompt
                    pdf_bytes = uploaded_file.read()
                    
                    model = genai.GenerativeModel('gemini-flash-latest')
                    
                    progress_bar.progress(30)
                    
                    response = model.generate_content([
                        {
                            "mime_type": "application/pdf",
                            "data": pdf_bytes
                        },
                        "Analyze the uploaded PDF meticulously. Extract all text, preserve the hierarchical structure (headings, body text), and identify tables. Reconstruct this content into a clean Markdown format."
                    ])
                    
                    progress_bar.progress(80)
                    
                    markdown_content = response.text
                    
                    if markdown_content:
                        docx_bytes = markdown_to_docx(markdown_content)
                        progress_bar.progress(100)
                        
                        st.success("Conversion successful!")
                        
                        # Preview
                        with st.expander("Preview Extracted Content"):
                            st.markdown(markdown_content)
                        
                        # Download button
                        st.download_button(
                            label="Download Word File",
                            data=docx_bytes,
                            file_name=uploaded_file.name.replace(".pdf", ".docx"),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("AI failed to extract content from the PDF.")
                        
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()
